from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT.parent / "Начален проект на дипломна работа.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Страница ")
    run.font.name = "Times New Roman"
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True

    caption = doc.styles.add_style("Надпис", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(11)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.0

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAD3")
        if widths:
            cell.width = Cm(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if widths:
                cells[idx].width = Cm(widths[idx])
    doc.add_paragraph()
    return table


def add_title_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    for line in (
        "[НАИМЕНОВАНИЕ НА ВИСШЕТО УЧИЛИЩЕ]",
        "[ФАКУЛТЕТ / ДЕПАРТАМЕНТ]",
        "[КАТЕДРА]",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).bold = True
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("НАЧАЛЕН ПРОЕКТ НА ДИПЛОМНА РАБОТА")
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("на тема")
    r.italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("„Модулна mesh система за наблюдение и автоматично поливане на стайни растения“")
    r.bold = True
    r.font.size = Pt(15)
    for _ in range(5):
        doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = ["Разработил:", "Факултетен №:", "Научен ръководител:"]
    values = ["[име и фамилия]", "[номер]", "[име, научна степен и длъжност]"]
    for i, (label, value) in enumerate(zip(labels, values)):
        set_cell_text(table.cell(i, 0), label, bold=True)
        set_cell_text(table.cell(i, 1), value)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[град], 2026 г.")
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "СЪДЪРЖАНИЕ (работна структура)", 1)
    lines = [
        "Анотация",
        "Използвани съкращения и означения",
        "Увод",
        "ГЛАВА ПЪРВА: Обзор на обекта на разработката",
        "ГЛАВА ВТОРА: Анализ на съществуващи технически решения",
        "ГЛАВА ТРЕТА: Проектиране и разработване на модулна mesh система за автоматично поливане",
        "ГЛАВА ЧЕТВЪРТА: Технико-икономическа обосновка и приложимост",
        "Заключение",
        "Използвана литература",
        "Приложения",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(line)
    add_body(doc, "Бележка: при оформяне на окончателната работа това работно съдържание следва да се замени с автоматично генерирано съдържание в Word.", first_line=False)


def add_document(doc):
    add_title_page(doc)
    add_contents(doc)

    add_heading(doc, "АНОТАЦИЯ", 1)
    add_body(doc, "Дипломната работа разглежда проектирането на модулна IoT система за наблюдение и автоматично поливане на стайни растения. Системата се състои от независими plant nodes, разположени при отделни саксии, и front/root node, който осигурява връзка между локалната mesh мрежа и сървърна система за наблюдение. Основният принцип е вземането на решение за поливане да остава локално в plant node, така че липсата на интернет или временна недостъпност на сървъра да не нарушава безопасната работа.")
    add_body(doc, "Към настоящия начален проект са определени базова хардуерна конфигурация, софтуерни отговорности, протоколен договор и план за последващо разработване и изпитване. Документът не представя окончателни измервания или готова софтуерна реализация; те ще бъдат добавени след калибрация на сензорите и лабораторни изпитвания.")
    add_heading(doc, "Ключови думи", 2)
    add_body(doc, "автоматично поливане, Internet of Things, IoT, mesh мрежа, ESP32-S3, plant node, front node, сензорно наблюдение, локална безопасност", first_line=False)

    add_heading(doc, "ИЗПОЛЗВАНИ СЪКРАЩЕНИЯ И ОЗНАЧЕНИЯ", 1)
    add_table(doc, ["Съкращение", "Значение"], [
        ("IoT", "Интернет на нещата"),
        ("ADC", "Аналогово-цифров преобразувател"),
        ("I²C", "Двупроводна серийна шина"),
        ("MOSFET", "Полеви транзистор с изолиран затвор"),
        ("REST", "Архитектурен стил за уеб интерфейс"),
        ("plant node", "Локален модул към една саксия"),
        ("front/root node", "Челен възел и мост към външна мрежа"),
    ], widths=[4.0, 11.0])

    add_heading(doc, "УВОД", 1)
    add_body(doc, "Поддържането на стайни растения изисква периодично наблюдение на влажността на почвата, наличието на вода и условията на средата. Ръчното поливане е лесно за единична саксия, но при повече растения или при отсъствие на потребителя носи риск от пресъхване, преполиване и недостатъчен контрол върху условията. Това създава предпоставки за разработване на модулна система, която измерва локални параметри, оценява ограниченията за безопасност и управлява поливането контролирано.")
    add_body(doc, "Обект на разработката е разпределена система за наблюдение и поливане на стайни растения. Предмет на разработката са хардуерната архитектура на автономен plant node, комуникацията в mesh мрежа, локалният алгоритъм за решение и ролята на front node като мост към HTTP/REST услуга.")
    add_heading(doc, "Цел и задачи", 2)
    add_body(doc, "Целта е да се проектира модулна и разширяема система, при която всеки plant node може да измерва състоянието на една саксия и да взема безопасно локално решение за поливане, а front node агрегира информацията за наблюдение и конфигуриране.")
    for task in [
        "Да се анализират съществуващи решения за автоматично поливане, сензори и комуникационни архитектури.",
        "Да се проектира хардуерна конфигурация на plant node с измерване на почвена влажност, осветеност, температура, влажност на въздуха, ниво на резервоара и заряд на батерията.",
        "Да се определи локален алгоритъм за поливане с ограничения за заряд, резервоар, сензорна достоверност, продължителност и дневно количество.",
        "Да се дефинира протокол за telemetry, watering event, конфигурация и командни съобщения без използване на MQTT.",
        "Да се разработи и изпита прототип, включително сценарии при прекъсната интернет връзка.",
    ]:
        add_bullet(doc, task)

    add_heading(doc, "ГЛАВА ПЪРВА: ОБЗОР НА ОБЕКТА НА РАЗРАБОТКАТА", 1)
    add_heading(doc, "1.1. Автоматизирани системи за напояване", 2)
    add_body(doc, "Автоматизираните системи за поливане обичайно използват времево управление, измерване на почвена влажност или комбинирано правило. При стайни растения измерването на действителната влажност е предпочитано пред фиксиран график, защото видът на растението, субстратът, температурата и осветеността влияят на нуждата от вода. В окончателната работа този раздел ще включи литературен обзор с коректно цитирани източници.")
    add_heading(doc, "1.2. Особености при отглеждане на стайни растения", 2)
    add_body(doc, "Различните растения имат различни допустими интервали на влажност, светлина и температура. Затова системата използва локален plant profile. Външни източници могат да предложат начални стойности, но критичното решение за поливане не зависи от външна услуга и винаги се ограничава от локалните правила за безопасност.")
    add_heading(doc, "1.3. Измерване на влажност на почвата", 2)
    add_body(doc, "За прототипа се предвижда капацитивен сензор за почвена влажност с аналогов изход. Тъй като стойностите зависят от сензора, почвата и напрежението на захранване, в практическата част ще се извърши калибрация за сухо и контролно навлажнено състояние. До тази калибрация всички прагове се разглеждат като работни placeholders.")
    add_heading(doc, "1.4. Измерване на осветеност, температура и влажност на въздуха", 2)
    add_body(doc, "Текущата базова конфигурация използва AHT10 за температура и относителна влажност на въздуха и TEMT6000 за аналогово измерване на осветеност. Сензорът AHT10 трябва да бъде разположен далеч от топлината на контролера и помпата, а фотосензорът трябва да има защитен, но свободен светлинен отвор.")
    add_heading(doc, "1.5. Безжични IoT системи и mesh архитектури", 2)
    add_body(doc, "Mesh архитектурата е избрана, за да позволи разширяване с няколко plant nodes и предаване през междинни възли при необходимост. Front node е единственият възел с планирана връзка към интернет. Plant nodes не изискват пряк достъп до интернет и продължават да работят автономно по локалния си профил.")
    add_heading(doc, "1.6. Изводи от глава първа", 2)
    add_body(doc, "Необходима е модулна архитектура, при която измерването и safety проверките са локални, комуникацията е отделена от хардуерните драйвери, а отдалечените команди не могат да заобиколят локалните ограничения.")

    add_heading(doc, "ГЛАВА ВТОРА: АНАЛИЗ НА СЪЩЕСТВУВАЩИ ТЕХНИЧЕСКИ РЕШЕНИЯ", 1)
    add_heading(doc, "2.1. Готови системи за автоматично поливане", 2)
    add_body(doc, "В окончателната версия ще бъдат сравнени готови потребителски решения според измервани параметри, локална автономност, възможност за калибрация, разширяемост и защити. Тук не се правят окончателни твърдения без проверени източници.")
    add_heading(doc, "2.2. Сравнение на микроконтролери", 2)
    add_body(doc, "ESP32-S3 Zero е избран като работна база за прототипа поради наличието на Wi-Fi функционалност, достатъчен изчислителен ресурс и подходящи интерфейси за сензори. Сравнението с други контролери ще се допълни с измерими критерии: консумация, налични ADC канали, библиотечна поддръжка, цена и възможности за безжична мрежа.")
    add_heading(doc, "2.3. Сензори и управление на помпа", 2)
    add_body(doc, "Изборът на капацитивен датчик за влажност, AHT10 и TEMT6000 е работен избор за прототипа. Помпата се управлява чрез нискостранен N-канален MOSFET IRLB3034 с резистор на gate, pull-down резистор, flyback диод и локален bulk кондензатор. В последващата работа ще се сравнят варианти на MOSFET-и и ще се проверят падове на напрежение, ток и нагряване.")
    add_heading(doc, "2.4. Анализ на комуникационни архитектури", 2)
    add_body(doc, "Не се използва MQTT. Локалната мрежа и HTTP/REST мостът са отделени: front node приема съобщения от plant nodes и предава telemetry и събития към backend чрез HTTP/REST. Точният транспортен механизъм за mesh остава предмет на следващ етап и трябва да бъде валидиран в лабораторен тест.")
    add_heading(doc, "2.5. Избор на концепция", 2)
    add_body(doc, "Избраната концепция комбинира локално решение за поливане, локално съхранение на plant profile и централизирано наблюдение. Това намалява зависимостта от интернет и не допуска backend или dashboard да включат помпата без приемане от локалния SafetyManager.")

    add_heading(doc, "ГЛАВА ТРЕТА: ПРОЕКТИРАНЕ И РАЗРАБОТВАНЕ НА МОДУЛНА MESH СИСТЕМА ЗА АВТОМАТИЧНО ПОЛИВАНЕ", 1)
    add_heading(doc, "3.1. Обща архитектура на системата", 2)
    add_body(doc, "Системата има три логически нива: plant node, front/root node и backend/dashboard. Plant node събира данни, пази профил и управлява помпата локално. Front node поддържа регистър на възлите, буферира съобщения при липса на интернет и в бъдеще предава данни към backend. Backend и dashboard са предназначени за наблюдение, история и изпращане на предложения за конфигурация или команда.")
    add_table(doc, ["Компонент", "Роля", "Състояние в началния проект"], [
        ("Plant node", "Локално измерване, safety проверка и поливане", "Дефинирани интерфейси"),
        ("Front/root node", "Mesh root, регистър, HTTP/REST мост", "Дефинирани интерфейси"),
        ("Backend", "Приемане на telemetry, история, конфигурации и команди", "Само резервирана структура"),
        ("Dashboard", "Наблюдение и редактиране на профили", "Само резервирана структура"),
    ], widths=[4.0, 7.0, 4.0])
    add_heading(doc, "3.2. Архитектура на plant node", 2)
    add_body(doc, "Plant node включва ESP32-S3 Zero, LiPo захранване, сензори, помпен драйвер и локално хранилище. Батерията се свързва към клемите B+/B- на зарядното, а товарът използва OUT+/OUT-, за да не се заобикаля защитната схема на модула. Конкретни пинове и електрически параметри ще се фиксират след проверка на физическата платка.")
    add_heading(doc, "3.3. Сензорна подсистема и изпълнителен механизъм", 2)
    add_body(doc, "Почвената влажност, зарядът на батерията и осветеността се очаква да бъдат измерени през ADC. AHT10 използва I²C. Поплавъковият датчик указва ниско ниво на вода. Помпата е 3,3 V вариант и се включва през IRLB3034 в нискостранна конфигурация. Реални драйвери не са част от началния проект.")
    add_heading(doc, "3.4. Софтуерна архитектура и състояния", 2)
    add_body(doc, "Предвиденият цикъл на plant node е: BOOT, LOAD_CONFIG, JOIN_MESH, SELF_TEST, MEASURE, SAFETY_CHECK, DECIDE, WATER_OR_SKIP, SETTLE, REPORT_TO_FRONT_NODE и SLEEP_OR_WAIT. Началното хранилище съдържа класови stubs за ConfigManager, SensorManager, PumpController, SafetyManager, IrrigationController, StateMachine, MeshClient и StorageManager. Тези класове още не изпълняват хардуерна функция.")
    add_heading(doc, "3.5. Алгоритъм за автоматично поливане", 2)
    add_body(doc, "Алгоритъмът първо проверява дали измерването на почвата е валидно, дали резервоарът не е празен и дали напрежението на батерията е над безопасния праг. След това проверява lockout периода, максималното количество за сесия и натрупаното дневно количество. Само ако всички условия са изпълнени и влажността е под минималния праг, може да бъде изчислена кратка продължителност на поливане според калибрирания дебит. След поливане се изчаква settle период и се отчита събитие.")
    add_heading(doc, "3.6. Защитни режими", 2)
    for rule in [
        "Помпата е изключена по подразбиране след reset.",
        "Поливането се забранява при нисък заряд на батерията.",
        "Поливането се забранява при празен резервоар.",
        "Поливането се забранява при невалидно измерване на почвата.",
        "Ограничават се продължителността на сесията и максималното количество вода за деня.",
        "След поливане се налага lockout период.",
        "Ръчна команда се оценява от същите локални правила и може да бъде отказана.",
    ]:
        add_bullet(doc, rule)
    add_heading(doc, "3.7. Комуникация и формат на съобщенията", 2)
    add_body(doc, "Дефинирани са логически типове telemetry, watering_event, config_set, command, command_result, heartbeat и error. Всяко съобщение има версия, идентификатор на устройство, пореден номер и времеви маркер. Примерните JSON структури и отговорностите на участниците са описани в repository файла docs/communication_protocol.md. Парсерът и транспортът са оставени като празни stubs.")
    add_heading(doc, "3.8. Данни за plant profile", 2)
    add_table(doc, ["Група", "Примерни полета"], [
        ("Идентификация", "device_id, plant_name, scientific_name, profile_source"),
        ("Прагове", "soil_moisture_min, soil_moisture_target, light_min_lux, air_temperature_min/max"),
        ("Поливане", "pump_ml_per_second, pulse_ml, max_session_ml, max_daily_ml"),
        ("Безопасност", "lockout_minutes, settle_minutes, low_battery_cutoff"),
        ("Периодичност", "reporting_interval_minutes, measurement_interval_minutes"),
    ], widths=[4.0, 11.0])
    add_body(doc, "Стойностите не са попълнени умишлено. Те трябва да се зададат след калибрация за конкретна почва, растение, сензор, помпа и резервоар.")
    add_heading(doc, "3.9. Състояние на софтуерната разработка", 2)
    add_body(doc, "Репозиториумът съдържа PlatformIO-ориентирана структура за plant node и front node, общи protocol definitions, README файлове и документация. Backend, база данни и dashboard са само резервирани структури. Не са реализирани mesh мрежа, драйвери, постоянно хранилище, HTTP клиент или визуализация.")

    add_heading(doc, "ГЛАВА ЧЕТВЪРТА: ТЕХНИКО-ИКОНОМИЧЕСКА ОБОСНОВКА И ПРИЛОЖИМОСТ", 1)
    add_heading(doc, "4.1. Обосновка на необходимостта", 2)
    add_body(doc, "Системата е приложима при поддържане на няколко растения в жилищни, офисни, лабораторни и учебни среди. Модулният подход позволява отделяне на профилите на растенията, локално отчитане на условията и поетапно разширяване без изграждане на централизирана кабелна инфраструктура.")
    add_heading(doc, "4.2. Разходи и измерими показатели", 2)
    add_body(doc, "В окончателната работа следва да се изготви таблица с реални цени на използваните модули, консумация на енергия, дебит на помпата и оценка на надеждността. В началния проект не се въвеждат предположени цени или ефекти, за да не се представят непроверени стойности като резултати.")
    add_heading(doc, "4.3. Мащабируемост и поддръжка", 2)
    add_body(doc, "Независимите plant nodes позволяват добавяне на нови саксии. Front node централизира наблюдението, без да поема безопасната логика на помпата. Това разделение улеснява диагностицирането: локален проблем със сензор или резервоар остава видим като telemetry/error събитие, без да блокира останалите възли.")
    add_heading(doc, "4.4. Изводи от глава четвърта", 2)
    add_body(doc, "Практическата стойност на системата зависи от надеждната калибрация и безопасната локална логика, а не само от наличието на дистанционно управление. Следващите етапи трябва да измерят именно тези показатели.")

    add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1)
    add_body(doc, "Разработен е начален проект на модулна система за наблюдение и автоматично поливане на стайни растения. Изяснени са отговорностите на plant node, front/root node и бъдещ backend, дефинирани са предпазните ограничения и е избрана работна хардуерна база. Създадена е начална структура на repository и договор за съобщенията, без да се претендира за завършена реализация. Следващата практическа задача е калибрация на сензорите и помпата, последвана от минимален тест на две устройства за discovery и telemetry.")

    add_heading(doc, "ИЗПОЛЗВАНА ЛИТЕРАТУРА (ЗА ПОПЪЛВАНЕ)", 1)
    add_body(doc, "В окончателната работа тук трябва да се добавят само реално използвани и проверени източници: научни публикации, техническа документация за ESP32-S3, AHT10, TEMT6000, MOSFET-а и зарядния модул, както и източници за автоматизирано поливане и безжични мрежи. Не са включени фиктивни библиографски записи.")

    add_heading(doc, "ПРИЛОЖЕНИЯ", 1)
    add_heading(doc, "Приложение А. План на следващите стъпки", 2)
    add_table(doc, ["Етап", "Резултат", "Критерий за приемане"], [
        ("Калибрация", "Прагове за влажност и дебит на помпата", "Документирани измервания"),
        ("Хардуерни драйвери", "Четене на всички сензори и безопасно OFF състояние", "Лабораторен тест"),
        ("Локална логика", "Safety и irrigation алгоритъм", "Тестове на откази"),
        ("Комуникация", "Директен тест между два възела", "Telemetry се приема коректно"),
        ("Backend и dashboard", "Наблюдение, история и конфигурация", "End-to-end демонстрация"),
    ], widths=[3.0, 7.0, 5.0])


def main():
    document = Document()
    configure_document(document)
    document.core_properties.title = "Начален проект на дипломна работа"
    document.core_properties.subject = "Модулна mesh система за наблюдение и автоматично поливане на стайни растения"
    document.core_properties.author = "[име и фамилия]"
    add_document(document)
    document.save(OUTPUT)
    print("Diploma document generated.")


if __name__ == "__main__":
    main()
