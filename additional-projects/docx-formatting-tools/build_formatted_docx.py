from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUT = Path("outputs/Predlojenie-komandirovka_BG_ICSQE-2026_formatted.docx")
FONT = "Times New Roman"
TOTAL_WIDTH = 10092  # DXA; exact A4 usable width with 16 mm left/right margins


def set_font(run, size=10.1, bold=False, italic=False, color="000000"):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=65, start=90, bottom=65, end=90):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **edges):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge, spec in edges.items():
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        el = borders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            borders.append(el)
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), spec.get("color", "D9D9D9"))


def set_table_geometry(table, widths, indent=0):
    assert sum(widths) == TOTAL_WIDTH
    table.autofit = False
    tblpr = table._tbl.tblPr
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(TOTAL_WIDTH))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def clear_cell(cell):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        p._p.remove(r._r)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_runs(p, parts, align=None):
    if align is not None:
        p.alignment = align
    for text, opts in parts:
        set_font(p.add_run(text), **opts)


def set_keep(p, keep_next=False, keep_lines=True):
    ppr = p._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def section_heading(doc, number, title):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TOTAL_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E7E6E6")
    set_cell_borders(cell, bottom={"color": "7F7F7F", "sz": 6})
    set_cell_margins(cell, top=55, bottom=55, start=100, end=100)
    p = clear_cell(cell)
    add_runs(p, [(f"{number}.  {title}", {"size": 9.8, "bold": True})])
    p.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0.5)


def key_value_table(doc, rows, label_width=3200):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        p1 = clear_cell(cells[0])
        add_runs(p1, [(label, {"size": 9.4, "bold": True, "color": "404040"})])
        p2 = clear_cell(cells[1])
        add_runs(p2, [(value, {"size": 10.0})])
        for cell in cells:
            set_cell_borders(cell, bottom={"color": "D9D9D9", "sz": 3})
            set_cell_margins(cell, top=75, bottom=75)
    set_table_geometry(table, [label_width, TOTAL_WIDTH - label_width])
    return table


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(6)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.1)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style._element.rPr.rFonts.set(qn("w:cs"), FONT)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    settings = doc.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    setting = OxmlElement("w:compatSetting")
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), "15")
    compat.append(setting)


def build():
    doc = Document()
    set_doc_defaults(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, [("БЪЛГАРСКА АКАДЕМИЯ НА НАУКИТЕ", {"size": 10.8, "bold": True})])
    set_keep(p, keep_next=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, [("ИНСТИТУТ ПО ЕЛЕКТРОНИКА", {"size": 10.8, "bold": True})])
    set_keep(p, keep_next=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    add_runs(p, [("ПРЕДЛОЖЕНИЕ ЗА КОМАНДИРОВКА В Р БЪЛГАРИЯ", {"size": 14.0, "bold": True})])
    set_keep(p, keep_next=True)

    section_heading(doc, 1, "КОМАНДИРОВАНО ЛИЦЕ")
    key_value_table(doc, [
        ("Трите имена, научни степени и звания", "Йоан Николаев Трендафилов, инженер"),
    ], label_width=3550)

    section_heading(doc, 2, "ДАННИ ЗА КОМАНДИРОВКАТА")
    key_value_table(doc, [
        ("До гр./с.", "Равда; обратно в София"),
        ("Срок и дати", "6 дни — от 20.09.2026 до 25.09.2026 г."),
        ("Маршрут", "София — Равда — София"),
        ("Мотивировка", 'Участие в 24-th International Conference and School on Quantum Electronics “Laser Physics and Applications”.'),
        ("Ще участва ли с доклад", "ДА"),
    ], label_width=2250)

    section_heading(doc, 3, "ФИНАНСОВИ УСЛОВИЯ")
    fin = doc.add_table(rows=1, cols=3)
    headers = ["", "Разход", "Осигуряване"]
    for cell, text in zip(fin.rows[0].cells, headers):
        set_cell_shading(cell, "F2F2F2")
        set_cell_borders(cell, bottom={"color": "A6A6A6", "sz": 5})
        p = clear_cell(cell)
        add_runs(p, [(text, {"size": 9.2, "bold": True, "color": "404040"})])
    set_repeat_table_header(fin.rows[0])
    fin_rows = [
        ("а)", "Пътни", "За сметка на договор КП-06-МНФ/13 от 18.05.2026 г."),
        ("б)", "Дневни", "За сметка на договор Д01-131 от 2025 г.; 22 евро на ден"),
        ("в)", "Квартирни", "За сметка на договор КП-06-МНФ/13 от 18.05.2026 г."),
        ("г)", "Такса за правоучастие", "ДА — 380 евро; за сметка на договор Д01-131 от 2025 г."),
        ("д)", "Други", "—"),
    ]
    for letter, kind, detail in fin_rows:
        cells = fin.add_row().cells
        for cell in cells:
            set_cell_borders(cell, bottom={"color": "D9D9D9", "sz": 3})
            set_cell_margins(cell, top=65, bottom=65)
        for cell, text, bold, align in (
            (cells[0], letter, True, WD_ALIGN_PARAGRAPH.CENTER),
            (cells[1], kind, True, WD_ALIGN_PARAGRAPH.LEFT),
            (cells[2], detail, False, WD_ALIGN_PARAGRAPH.LEFT),
        ):
            p = clear_cell(cell)
            add_runs(p, [(text, {"size": 9.5, "bold": bold})], align=align)
    set_table_geometry(fin, [550, 2450, 7092])

    section_heading(doc, 4, "ОТЧЕТНОСТ")
    key_value_table(doc, [
        ("Командировки през настоящата година", "НЕ"),
        ("Представени отчетни доклади през отчетната година", "НЕ"),
    ], label_width=5000)

    section_heading(doc, 5, "ПРЕДЛОЖЕНИЕ И СЪГЛАСУВАНЕ")
    approvals = doc.add_table(rows=0, cols=2)
    approval_rows = [
        ("Ръководител на договор", "доц. Таня Драйшу"),
        ("Ръководител на лаборатория", "Лаб. „Лазерна локация“, доц. Захари Пешев"),
        ("Главен счетоводител на ИЕ-БАН", "К. Родева"),
    ]
    for role, name in approval_rows:
        cells = approvals.add_row().cells
        p1 = clear_cell(cells[0])
        add_runs(p1, [(role, {"size": 9.2, "bold": True, "color": "404040"})])
        p2 = clear_cell(cells[1])
        add_runs(p2, [(name, {"size": 9.6, "italic": True})])
        for cell in cells:
            set_cell_borders(cell, bottom={"color": "D9D9D9", "sz": 3})
            set_cell_margins(cell, top=62, bottom=62)
    set_table_geometry(approvals, [3450, 6642])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(3)
    add_runs(note, [("Бележка: при пътуване за сметка на източници, външни на ИЕ-БАН, командировката се съгласува само с ръководителя на лабораторията.", {"size": 8.3, "italic": True, "color": "595959"})])
    set_keep(note, keep_next=True)

    sig = doc.add_table(rows=2, cols=2)
    sig_labels = [
        "ПОДПИС НА КОМАНДИРОВАНИЯ",
        "РЪКОВОДИТЕЛ НА ЛАБОРАТОРИЯ",
        "РЪКОВОДИТЕЛ НА ДОГОВОР",
        "ГЛАВЕН СЧЕТОВОДИТЕЛ",
    ]
    for cell, label in zip([c for row in sig.rows for c in row.cells], sig_labels):
        set_cell_borders(cell, top={"color": "FFFFFF", "sz": 0})
        set_cell_margins(cell, top=100, start=90, bottom=55, end=150)
        p = clear_cell(cell)
        p.paragraph_format.space_after = Pt(13)
        add_runs(p, [(label, {"size": 8.8, "bold": True})])
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        add_runs(p2, [("подпис: ______________________________", {"size": 8.7, "color": "595959"})])
    set_table_geometry(sig, [5046, 5046])

    # Prevent table rows from splitting across pages.
    for table in doc.tables:
        for row in table.rows:
            trpr = row._tr.get_or_add_trPr()
            trpr.append(OxmlElement("w:cantSplit"))

    doc.core_properties.title = "Предложение за командировка в Р България — ICSQE 2026"
    doc.core_properties.subject = "Форматирано предложение за командировка"
    doc.core_properties.author = "Институт по електроника — БАН"
    doc.core_properties.keywords = "командировка, ICSQE 2026, Равда"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
