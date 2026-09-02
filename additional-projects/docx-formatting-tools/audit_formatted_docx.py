from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PATH = Path("outputs/Predlojenie-komandirovka_BG_ICSQE-2026_formatted.docx")
doc = Document(PATH)

required = [
    "Йоан Николаев Трендафилов",
    "20.09.2026",
    "25.09.2026",
    "София — Равда — София",
    "24-th International Conference and School on Quantum Electronics",
    "КП-06-МНФ/13",
    "Д01-131",
    "380 евро",
    "доц. Таня Драйшу",
    "доц. Захари Пешев",
    "К. Родева",
]

all_text = []
for p in doc.paragraphs:
    all_text.append(p.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text)
joined = "\n".join(all_text)

section = doc.sections[0]
usable_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
usable_height = section.page_height.twips - section.top_margin.twips - section.bottom_margin.twips

table_reports = []
for index, table in enumerate(doc.tables, 1):
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    grid = [int(c.get(qn("w:w"))) for c in table._tbl.tblGrid]
    cell_widths_ok = True
    for row in table.rows:
        values = []
        for cell in row.cells:
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            values.append(int(tcw.get(qn("w:w"))))
        if values != grid:
            cell_widths_ok = False
    table_reports.append({
        "table": index,
        "rows": len(table.rows),
        "cols": len(grid),
        "tblW": int(tblw.get(qn("w:w"))),
        "grid_sum": sum(grid),
        "cell_widths_match_grid": cell_widths_ok,
    })

with zipfile.ZipFile(PATH) as zf:
    names = set(zf.namelist())
    document_xml = zf.read("word/document.xml")

report = {
    "opens_with_python_docx": True,
    "sections": len(doc.sections),
    "page": {
        "width_twips": section.page_width.twips,
        "height_twips": section.page_height.twips,
        "usable_width_twips": usable_width,
        "usable_height_twips": usable_height,
    },
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "required_text": {text: text in joined for text in required},
    "table_geometry": table_reports,
    "has_section_properties": b"<w:sectPr" in document_xml,
    "has_core_properties": "docProps/core.xml" in names,
    "has_settings": "word/settings.xml" in names,
}

assert all(report["required_text"].values())
assert report["has_section_properties"]
assert report["has_core_properties"]
assert report["has_settings"]
assert all(t["tblW"] == usable_width == t["grid_sum"] for t in table_reports)
assert all(t["cell_widths_match_grid"] for t in table_reports)
print(json.dumps(report, ensure_ascii=False, indent=2))
